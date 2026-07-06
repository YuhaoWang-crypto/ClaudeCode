"""
make_reports.py  --  build two standalone .docx deliverables:
  Report 1: open-source reproduction of the RSV / C13H10N4S DFT project
  Report 2: computational screening for electrochemical sensors (DA/AA/UA)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

HERE = os.path.dirname(os.path.abspath(__file__))
BASE = os.path.join(HERE, "..")
FIG = os.path.join(BASE, "figures")
OUT = os.path.join(BASE, "reports")
os.makedirs(OUT, exist_ok=True)

ACCENT = RGBColor(0xC0, 0x5A, 0x1A)

def new_doc(title, subtitle):
    doc = Document()
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(10.5)
    # CJK font
    from docx.oxml.ns import qn
    st.element.rPr.rFonts.set(qn("w:eastAsia"), "SimSun")
    h = doc.add_paragraph(); h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = h.add_run(title); r.bold = True; r.font.size = Pt(18); r.font.color.rgb = ACCENT
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run(subtitle); rs.italic = True; rs.font.size = Pt(11)
    doc.add_paragraph()
    return doc

def H(doc, text, lvl=1):
    p = doc.add_heading(level=lvl)
    r = p.add_run(text); r.font.color.rgb = ACCENT if lvl == 1 else RGBColor(0x33,0x33,0x33)
    return p

def P(doc, text, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = bold
    return p

def BUL(doc, text):
    doc.add_paragraph(text, style="List Bullet")

def TBL(doc, header, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(header)); t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, htext in enumerate(header):
        c = t.rows[0].cells[i]; c.text = ""
        run = c.paragraphs[0].add_run(htext); run.bold = True; run.font.size = Pt(9.5)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""; run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
    return t

def FIGURE(doc, path, caption, width=6.2):
    if os.path.exists(path):
        doc.add_picture(path, width=Inches(width))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        c = doc.add_paragraph(); c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rc = c.add_run(caption); rc.italic = True; rc.font.size = Pt(9)


# ============================================================ REPORT 1 =========
def report1():
    doc = new_doc(
        "开源模型计算复现报告",
        "基于 fairchem UMA 通用机器学习势与 PySCF 的 RSV / C₁₃H₁₀N₄S 项目复现"
        "（对应原项目 CCB012626YW01）")

    H(doc, "1. 项目概述")
    P(doc, "本报告使用开源工具复现原 DFT 计算项目中可由能量/受力与电子结构获得的交付内容。"
           "项目包含两个独立分子体系：")
    BUL(doc, "RSV = 罗苏伐他汀 (Rosuvastatin, C₂₂H₂₈FN₃O₆S)，与 n 型 TiO₂ 和 p 型 TiO "
             "氧化物团簇结合，考察静态络合常数 Kₐ 与结合自由能 ΔG_bind。")
    BUL(doc, "C₁₃H₁₀N₄S = 3-巯基-4-苯基-5-(吡啶基)-1,2,4-三氮唑（缓蚀剂），吸附在 Fe(BCC)(110) "
             "表面，考察吸附能、差分电荷密度、Bader 电荷与前线分子轨道 (FMO)。")

    H(doc, "2. 计算方法")
    BUL(doc, "机器学习势：fairchem-core，检查点 uma-s-1p1（facebook/UMA），CPU 运行；"
             "oc20 任务用于金属表面吸附，omol 任务用于分子/团簇结合。几何优化用 ASE LBFGS，"
             "收敛判据 F_max = 0.05 eV/Å。")
    BUL(doc, "量子化学：PySCF，B3LYP/6-31G(d) 用于前线分子轨道；PBE 用于电子密度。")
    BUL(doc, "电荷分析：Henkelman 课题组 bader 程序（与 VASP 配套的标准工具）。")
    P(doc, "关键说明：机器学习原子间势 (MLIP) 只输出总能量与原子受力，不产生波函数或电子密度。"
           "因此凡涉及电子密度/轨道的量（差分电荷密度、Bader 电荷、分子轨道）必须由真正的 DFT "
           "给出，本项目用 PySCF/QE 补齐。", bold=False)

    H(doc, "3. 可复现性映射")
    TBL(doc, ["交付内容", "原方法", "MLIP 可否", "本项目实现"],
        [["C₁₃H₁₀N₄S/Fe(110) 吸附能", "VASP", "✅ 可", "UMA oc20"],
         ["RSV–TiOₓ 结合能/ΔG/Kₐ", "Gaussian", "✅ 可", "UMA omol + 谐振 ΔG"],
         ["几何优化（全部）", "两者", "✅ 可", "UMA 弛豫"],
         ["差分电荷密度 Δρ", "VASP", "❌ 力场不可", "PySCF/QE 密度（已跑通）"],
         ["Bader 电荷", "VASP", "❌ 力场不可", "PySCF + bader（已跑通）"],
         ["前线分子轨道 FMO", "Gaussian", "❌ 力场不可", "PySCF B3LYP/6-31G*"]])

    H(doc, "4. 计算结果")
    H(doc, "4.1 C₁₃H₁₀N₄S 在 Fe(110) 表面的吸附能", 2)
    P(doc, "分子通过 N/S 孤对原子化学吸附；平躺（π 平行）构型不稳定会脱附。四种对接构型弛豫结果：")
    TBL(doc, ["构型", "最近 S/N–Fe 距离", "吸附能 Eads"],
        [["N 端-顶位（最优）", "1.93 Å", "−1.80 eV"],
         ["S 端-顶位", "2.23 Å", "−1.04 eV"],
         ["N 端-桥位", "1.96 Å", "−0.47 eV"],
         ["平躺/π 平行", "脱附 (4.4 Å)", "+0.93 eV"]])
    P(doc, "本项目 UMA 结果 −1.80 eV，原报告 VASP 结果 −2.88 eV：符号与量级一致，均为化学吸附，"
           "结合原子正确。约 1 eV 的差异来自泛函不同（UMA 的 oc20 头偏向 RPBE，结合弱于原报告的 PBE），"
           "以及原报告可能采用了不同（多齿）结合模式。")

    H(doc, "4.2 RSV 与 TiOₓ 团簇的结合能、ΔG_bind 与 Kₐ", 2)
    TBL(doc, ["体系", "原报告 ΔG_bind", "原报告 Kₐ", "UMA ΔE_bind", "UMA ΔG_bind(谐振)", "UMA log₁₀Kₐ"],
        [["n 型 TiO₂–RSV", "−1.808 eV", "3.8×10³⁰", "−1.12 eV", "−1.28 eV", "≈ 22"],
         ["p 型 TiO–RSV", "−3.372 eV", "1.0×10⁵⁷", "−5.58 eV", "−5.62 eV", "≈ 95"],
         ["相对 (p−n)", "p 更强", "p ≫ n", "−4.46 eV", "−4.34 eV", "p ≫ n ✓"]])
    P(doc, "原报告最强调的物理结论——p 型 TiO 对 RSV 的结合远强于 n 型 TiO₂，且两者 Kₐ 都是天文数字、"
           "只有相对大小有意义——被干净地复现。绝对值差异源于 UMA 的 omol 头训练泛函为 ωB97M-V（原报告为 "
           "B3LYP），以及 H 饱和团簇模型为本项目自建（原报告未给出确切团簇）。")

    H(doc, "4.3 C₁₃H₁₀N₄S 的前线分子轨道 (FMO)", 2)
    P(doc, "PySCF B3LYP/6-31G(d)，在 UMA 优化的气相几何上计算（自由分子）：")
    TBL(doc, ["描述符", "数值", "缓蚀剂意义"],
        [["E(HOMO)", "−6.32 eV", "向金属供电子能力"],
         ["E(LUMO)", "−1.30 eV", "接受电子（反馈）能力"],
         ["HOMO–LUMO 带隙", "5.02 eV", "动力学稳定性/反应活性"],
         ["化学硬度 η", "2.51 eV", "软硬程度"],
         ["偶极矩", "5.18 D", "极性/物理吸附倾向"]])
    P(doc, "说明：原报告表中 0.065 eV 的带隙是分子吸附在 Fe 金属上时的值（由金属态主导，非分子本征），"
           "与自由分子 HOMO/LUMO 是两个不同的量。上表给出的是可迁移的自由分子前线轨道描述符。")

    H(doc, "4.4 差分电荷密度与 Bader 电荷", 2)
    P(doc, "MLIP 无法给出这两项（无电子密度），但开源 DFT 可以。完整流程为：PySCF PBE 电子密度 → 同网格 "
           "cube → Δρ = ρ(复合物) − ρ(团簇) − ρ(吸附质) → Henkelman bader 分区 → 净电荷转移。")
    P(doc, "受本次纯 CPU 环境限制，真实 86–128 原子体系的自洽 DFT 无法算完，故在可行的闭壳层类比体系 "
           "（甲酸 HCOOH——模拟 RSV 羧基/羟基结合基团——吸附在 Ti₂O₄ 团簇）上端到端验证了整条流水线，"
           "并提供全尺寸驱动脚本与 Quantum ESPRESSO 输入文件供 GPU/集群运行。")
    TBL(doc, ["量", "数值"],
        [["∫Δρ dV（自洽性检验）", "0.000 e ✓"],
         ["Bader 总电子数", "100.58（100 的网格离散）"],
         ["净电荷转移", "0.21 e，吸附质 → 氧化物团簇"]])
    FIGURE(doc, os.path.join(FIG, "demo_delta_rho_slice.png"),
           "图 1. 差分电荷密度 Δρ（HCOOH/Ti₂O₄ 演示）：蓝=电子聚集，红=电子耗散。"
           "Ti–O 结合区极化清晰，对应原报告 Figure 1–3 的同类结果。")
    P(doc, "净电荷转移 0.21 e 与原报告的 0.10 e 属同一量级（原报告为 Fe→分子，此处为有机酸→氧化物，"
           "方向因化学不同而相反）。")

    H(doc, "5. 验证要点与局限")
    BUL(doc, "参考态一致性：首次 Fe(110) 计算给出 −10.8 eV 的错误值，查明是裸板参考态仅弛豫 6 步、"
             "高于真实表面弛豫最小值 11.7 eV 且分子脱附所致；修正参考态并重新多取向对接后得 −1.80 eV。")
    BUL(doc, "泛函差异：UMA 各任务头训练泛函（oc20≈RPBE，omol≈ωB97M-V）与原报告（PBE/B3LYP）不同，"
             "绝对值有系统偏移，符号/量级/相对趋势可迁移。")
    BUL(doc, "团簇模型近似、谐振 ΔG 的 MLIP-Hessian 噪声、Bader 网格离散（100.58 vs 100）等已在正文标注。")

    H(doc, "6. 结论")
    TBL(doc, ["交付内容", "状态"],
        [["几何优化（全部体系）", "✅ 已复现（UMA）"],
         ["Fe(110) 吸附能", "✅ 已复现（−1.80 vs −2.88 eV）"],
         ["RSV–TiOₓ 结合能与相对趋势", "✅ 已复现（p ≫ n）"],
         ["ΔG_bind、Kₐ", "✅ 谐振近似复现"],
         ["FMO / HOMO–LUMO", "✅ 已复现（PySCF）"],
         ["差分电荷密度", "✅ 流程已验证（PySCF/QE），全尺寸需 HPC"],
         ["Bader 电荷", "✅ 流程已验证（PySCF+bader），全尺寸需 HPC"]])
    P(doc, "全部 7 项交付内容均已具备开源复现路径：能量类由 fairchem UMA 直接给出，电子结构类由 "
           "PySCF/QE + bader 给出。唯一限制是全尺寸真实体系的 DFT 需比 4 核 CPU 更大的算力，"
           "驱动与输入文件均已备好。")

    out = os.path.join(OUT, "报告1_开源模型计算复现报告.docx")
    doc.save(out); print("wrote", out)


# ============================================================ REPORT 2 =========
def report2():
    doc = new_doc(
        "电化学传感器计算筛选报告",
        "基于第一性原理与机器学习势的传感描述符预测 —— 以多巴胺/抗坏血酸/尿酸为例")

    H(doc, "1. 背景与动机")
    P(doc, "对电化学传感器而言，真正有意义的量不是结合常数 Kₐ（它只说明分析物能否吸附），而是"
           "吸附/氧化后可测量的电子响应：电荷转移、功函数变化、态密度 (DOS) 变化、氧化还原电势。"
           "本报告将这些信号逐一对应到具体的开源计算方法，并在经典的多巴胺 (DA)/抗坏血酸 (AA)/"
           "尿酸 (UA) 传感体系上演示分子级描述符的预测。三者氧化峰在裸电极上重叠，好的传感材料"
           "必须把它们分开——正是描述符预测要解决的问题。")

    H(doc, "2. 传感信号 ↔ 描述符 ↔ 计算方法")
    TBL(doc, ["传感器模态（读什么）", "关键描述符", "计算引擎"],
        [["伏安型 (CV/DPV) 峰电位", "氧化还原电势 E°(vs SHE)；HOMO/IP、LUMO/EA", "PySCF + 隐式溶剂"],
         ["化学电阻 / FET 电导", "吸附电荷转移 ΔQ (Bader)、掺杂符号", "UMA 筛 + QE/GPAW + bader"],
         ["功函数型 (Kelvin/FET)", "功函数变化 Δφ", "QE/GPAW slab"],
         ["电子态响应", "DOS/PDOS 移动、费米能级移动", "QE/GPAW slab PDOS"],
         ["选择性/亲和", "吸附能 E_ads（目标 vs 干扰）", "UMA（可大规模快筛）"],
         ["分子固有反应性", "χ, η, ω, 偶极", "PySCF"]])

    H(doc, "3. 筛选工作流")
    BUL(doc, "① UMA 快筛：每个（分析物×电极/识别材料）对，优化 + 吸附能，按选择性排序。可上百分子。")
    BUL(doc, "② DFT 电子指纹：对命中做 PySCF（HOMO/LUMO、IP/EA、氧化还原电势+溶剂）"
             "与 QE/GPAW（ΔQ、Δφ、DOS 移动）。")
    BUL(doc, "③ 指纹表：每个分析物 → 描述符向量；比较正交性，判断哪种材料能把目标与干扰分开。")
    BUL(doc, "④（可选）把描述符当虚拟传感器阵列，或 ML 回归到实测响应（电子鼻式模式识别）。")

    H(doc, "4. 氧化还原电势的计算")
    P(doc, "一电子氧化 A → A⁺ + e⁻（作为反应性代理）：", bold=True)
    P(doc, "ΔG_ox = G(A⁺, 溶剂) − G(A, 溶剂)；  E_ox(vs SHE) = ΔG_ox − 4.44 V。")
    P(doc, "质子耦合 (PCET) 一氢一电子氧化（真实机理，抗氧化剂路径）：", bold=True)
    P(doc, "RedH(aq) → Red•(aq) + H⁺(aq) + e⁻；")
    P(doc, "E°(vs SHE, pH0) = [E(Red•,aq) − E(RedH,aq) + G(H⁺,aq)] − E_abs(SHE)；"
           "E°(pH7) = E°(pH0) − 0.05916 × 7（Nernst，1H⁺/1e⁻）。")

    H(doc, "5. 演示：DA / AA / UA")
    H(doc, "5.1 分子描述符指纹", 2)
    P(doc, "UMA 优化几何 + PySCF B3LYP/6-31G*，ddCOSMO 水：")
    TBL(doc, ["分析物", "HOMO(eV)", "带隙(eV)", "η(eV)", "ω(eV)", "μ(D)", "IP(eV)"],
        [["多巴胺", "−5.45", "5.74", "2.87", "1.16", "2.05", "7.21"],
         ["尿酸", "−5.90", "5.10", "2.55", "2.20", "3.02", "7.68"],
         ["抗坏血酸", "−6.04", "5.69", "2.84", "1.79", "7.19", "7.84"]])
    FIGURE(doc, os.path.join(FIG, "sensor_fingerprint.png"),
           "图 1. 左：预测伏安峰（朴素 1e⁻）；右：描述符指纹热图。每个分子在描述符空间中位置不同，"
           "这是传感材料能区分它们的物理依据。")

    H(doc, "5.2 质子耦合 (PCET) 氧化电势——对上实验", 2)
    P(doc, "朴素 1e⁻ 路径把 DA/AA/UA 的顺序算错（抗坏血酸被算成最难氧化），因为其真实氧化是质子耦合的。"
           "改用 1H⁺/1e⁻ PCET 中性自由基路径（对每个分子枚举所有 O–H/N–H，取最稳定自由基）后：")
    TBL(doc, ["分析物", "朴素 1e⁻（错）", "PCET, pH7", "实验, pH7"],
        [["抗坏血酸", "+1.51 V（最难 ✗）", "+0.38 V（最易 ✓）", "≈ 0.06–0.35 V"],
         ["多巴胺", "+0.94 V", "+0.46 V", "≈ 0.38 V"],
         ["尿酸", "+1.05 V", "+0.66 V", "≈ 0.59 V"]])
    FIGURE(doc, os.path.join(FIG, "sensor_pcet.png"),
           "图 2. 左：PCET 预测伏安图（实线=计算峰，虚线=实验）；右：朴素 vs PCET vs 实验。"
           "PCET 使顺序 AA<DA<UA 与实验一致，绝对值落在 ±0.1 V 内。")
    P(doc, "预测顺序 AA < DA < UA 与实验完全一致，绝对值在实验 ±0.1 V 内。物理上：最容易给出氢原子"
           "（最强抗氧化剂=抗坏血酸）的分子氧化电位最低——只有把耦合的质子算进去才能抓住。")

    H(doc, "6. 结论与对药物检测的意义")
    BUL(doc, "整条链完整：UMA 快筛吸附能/结合能 → PySCF 电子指纹 + PCET 氧化电势 → QE/GPAW 表面 "
             "ΔQ/Δφ/DOS，可直接套用于目标药物 + 电极材料的选材筛选。")
    BUL(doc, "关键教训：电化学描述符必须用对机理（PCET、pH）才能定量；朴素电子移除只能作粗略反应性代理。")
    BUL(doc, "应用方式：把分析物换成目标药物 + 干扰物的 SMILES，即可批量算指纹表，挑出信号分得开、"
             "选择性好的组合，再补表面描述符——即计算辅助的药物检测传感材料筛选。")

    H(doc, "7. 局限")
    BUL(doc, "氧化还原电势绝对值受 G(H⁺,aq)、E_abs(SHE) 参考值（用 4.28 V 会整体平移 +0.16 V，不改顺序）、"
             "泛函/基组、以及未计 ZPE/热校正影响；顺序与 ~0.1 V 精度为可靠交付。")
    BUL(doc, "表面描述符（ΔQ、Δφ、DOS）的全尺寸计算需 GPU/集群，驱动脚本已提供。")

    out = os.path.join(OUT, "报告2_电化学传感器计算筛选报告.docx")
    doc.save(out); print("wrote", out)


if __name__ == "__main__":
    report1()
    report2()
